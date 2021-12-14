from docx import Document
from paragraphs_extractor.file_iterator_interface import FileIteratorInterface

class DOCXIterator(FileIteratorInterface):
    def __init__(self, filename):
        super().__init__()
        self.filename = filename
        
        doc = Document(filename)
        for para in doc.paragraphs:
            if hasattr(para, 'text'):
                cleaned_text = para.text.replace('\n', '')
                if cleaned_text:
                    self.paragraphs.append(cleaned_text)